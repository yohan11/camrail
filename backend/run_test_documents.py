import sys
import os
import io

# Add the current directory to sys.path
sys.path.append(os.path.abspath(os.path.dirname(__file__)))

from fastapi.testclient import TestClient
from fpdf import FPDF
import docx

from app.main import app, startup_event
from app.database import SessionLocal
from app.models.schemas import User, Document, DocumentPage, AuditEvent

def generate_test_pdf() -> bytes:
    """Generates an in-memory 2-page PDF with text on page 1 and a table on page 2."""
    pdf = FPDF()
    # Page 1
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text="CAMRAIL Document d'Exploitation Ferroviaire - Reglement General.", new_x="LMARGIN", new_y="NEXT")
    pdf.multi_cell(0, 8, text="Ce document decrit l'ensemble des directives de securite et de conformite operationnelle pour les agents ferroviaires de CAMRAIL.")
    
    # Page 2
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.cell(0, 10, text="Tableau des Lignes et Stations :", new_x="LMARGIN", new_y="NEXT")
    with pdf.table() as table:
        table.row(["Station", "Ligne", "Statut"])
        table.row(["Douala Port", "Transcam 1", "Actif"])
        table.row(["Yaounde Voyageurs", "Transcam 1", "Actif"])
    
    return bytes(pdf.output())


def generate_test_docx() -> bytes:
    """Generates an in-memory DOCX with a paragraph and a table."""
    doc = docx.Document()
    doc.add_heading("CAMRAIL Manuel Technique de Maintenance", level=1)
    doc.add_paragraph("Consignes operationnelles pour la revision systematique des motrices et wagons.")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Composant"
    table.cell(0, 1).text = "Intervalle de Controle"
    table.cell(1, 0).text = "Essieux Moteurs"
    table.cell(1, 1).text = "Tous les 5000 km"
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def run_tests():
    with TestClient(app) as client:
        print("=" * 70)
        print("      RAILMIND LITE - DAY 2 DOCUMENT INGESTION & EXTRACTION")
        print("=" * 70)
        
        # 1. Startup & database check
        print("\n[Init] Initializing database and verifying users (via TestClient)...")
    
        # 1. Login as document_admin
        print("\n[Step 1] Authenticating as document_admin...")
        docadmin_login = {
            "username": "docadmin@camrail.net",
            "password": "docadminpassword"
        }
        res = client.post("/auth/login", data=docadmin_login)
        assert res.status_code == 200, f"Login failed: {res.text}"
        docadmin_token = res.json()["access_token"]
        docadmin_headers = {"Authorization": f"Bearer {docadmin_token}"}
        print(f"-> Login successful! Role: {res.json()['role']}")
        
        # Clean previous test documents if any
        db = SessionLocal()
        prev_docs = db.query(Document).filter(Document.title.in_(["Test Policy", "DOCX Maintenance Policy"])).all()
        for d in prev_docs:
            if d.file_url and os.path.exists(d.file_url):
                try:
                    os.remove(d.file_url)
                except Exception:
                    pass
            db.delete(d)
        db.commit()
        db.close()
        
        # 2. Generate 2-page PDF
        print("\n[Step 2] Generating 2-page test PDF (Text + Table)...")
        pdf_bytes = generate_test_pdf()
        print(f"-> PDF generated in-memory ({len(pdf_bytes)} bytes)")
        
        # 3. POST /documents with PDF
        print("\n[Step 3] Uploading PDF document (POST /documents)...")
        pdf_file_tuple = ("test_policy.pdf", pdf_bytes, "application/pdf")
        pdf_form_data = {
            "title": "Test Policy",
            "category": "Personnel Policy",
            "department": "Operations",
            "version": "1.0"
        }
        res = client.post("/documents", data=pdf_form_data, files={"file": pdf_file_tuple}, headers=docadmin_headers)
        assert res.status_code == 200, f"PDF upload failed: {res.text}"
        pdf_doc = res.json()
        pdf_doc_id = pdf_doc["id"]
        print(f"-> PDF Document uploaded successfully (ID: {pdf_doc_id})")
        print(f"   * Status: {pdf_doc['status']}")
        print(f"   * Checksum: {pdf_doc['checksum']}")
        assert pdf_doc["status"] == "indexed", f"Expected status 'indexed', got '{pdf_doc['status']}'"
        
        # 4. GET /documents/{id} detail & check pages / table extraction
        print(f"\n[Step 4] Fetching document details (GET /documents/{pdf_doc_id})...")
        res = client.get(f"/documents/{pdf_doc_id}", headers=docadmin_headers)
        assert res.status_code == 200, f"Get detail failed: {res.text}"
        detail = res.json()
        pages = detail.get("pages", [])
        pages.sort(key=lambda p: p["page_number"])
        print(f"-> Document retrieved with {len(pages)} extracted page(s)")
        assert len(pages) == 2, f"Expected 2 pages for PDF, got {len(pages)}"
        
        page1 = pages[0]
        page2 = pages[1]
        print(f"   * Page 1 (method={page1['extraction_method']}): {page1['extracted_text'][:60]}...")
        print(f"   * Page 2 (method={page2['extraction_method']}): {page2['extracted_text'][:60]}...")
        assert "|" in page2["extracted_text"], f"Table markdown '|' missing in Page 2 text: {page2['extracted_text']}"
        print("-> Table correctly extracted as Markdown formatted table on Page 2!")
        
        # 5. Re-upload identical PDF -> verify 409 Conflict
        print("\n[Step 5] Testing deduplication with duplicate PDF upload...")
        res = client.post("/documents", data=pdf_form_data, files={"file": pdf_file_tuple}, headers=docadmin_headers)
        assert res.status_code == 409, f"Expected 409 Conflict, got {res.status_code}: {res.text}"
        print(f"-> Duplicate correctly rejected with 409 Conflict: {res.json()['detail']}")
        
        # 6. Generate & upload DOCX
        print("\n[Step 6] Generating & uploading DOCX document (Text + Table)...")
        docx_bytes = generate_test_docx()
        docx_file_tuple = ("test_maintenance.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        docx_form_data = {
            "title": "DOCX Maintenance Policy",
            "category": "Maintenance",
            "department": "Technical",
            "version": "1.0"
        }
        res = client.post("/documents", data=docx_form_data, files={"file": docx_file_tuple}, headers=docadmin_headers)
        assert res.status_code == 200, f"DOCX upload failed: {res.text}"
        docx_doc = res.json()
        docx_doc_id = docx_doc["id"]
        print(f"-> DOCX Document uploaded (ID: {docx_doc_id}, Status: {docx_doc['status']})")
        
        # Verify DOCX single page extraction
        res = client.get(f"/documents/{docx_doc_id}", headers=docadmin_headers)
        docx_detail = res.json()
        docx_pages = docx_detail.get("pages", [])
        assert len(docx_pages) == 1, f"Expected 1 logical page for DOCX, got {len(docx_pages)}"
        print(f"-> DOCX successfully processed into 1 logical page (limitation handled correctly)")
        assert "|" in docx_pages[0]["extracted_text"], "Table markdown missing in DOCX page"
        
        # 7. Non-admin / Read-only permission check
        print("\n[Step 7] Testing permission restriction with read_only user...")
        readonly_login = {
            "username": "readonly@camrail.net",
            "password": "readonlypassword"
        }
        res = client.post("/auth/login", data=readonly_login)
        assert res.status_code == 200, f"Readonly login failed: {res.text}"
        readonly_token = res.json()["access_token"]
        readonly_headers = {"Authorization": f"Bearer {readonly_token}"}
        
        # Attempt upload
        unauth_file = ("unauth.pdf", b"%PDF-1.4 dummy", "application/pdf")
        res = client.post("/documents", data={"title": "Unauthorized", "category": "General", "department": "HQ"}, files={"file": unauth_file}, headers=readonly_headers)
        assert res.status_code == 403, f"Expected 403 Forbidden, got {res.status_code}"
        print(f"-> Upload correctly denied for read_only user (403 Forbidden): {res.json()['detail']}")
        
        # 8. Activate document
        print(f"\n[Step 8] Activating document (POST /documents/{pdf_doc_id}/activate)...")
        res = client.post(f"/documents/{pdf_doc_id}/activate", headers=docadmin_headers)
        assert res.status_code == 200, f"Activate failed: {res.text}"
        activated_doc = res.json()
        print(f"-> Document activated successfully! New status: {activated_doc['status']}")
        assert activated_doc["status"] == "active", f"Expected 'active', got '{activated_doc['status']}'"
        
        # 9. Verify Audit events recorded
        print("\n[Step 9] Verifying Audit Trail in Database...")
        db = SessionLocal()
        audit_events = db.query(AuditEvent).filter(AuditEvent.entity_type == "document").all()
        print(f"-> Total document audit events logged: {len(audit_events)}")
        for ev in audit_events:
            print(f"   * Action: {ev.action} | Entity ID: {ev.entity_id} | Actor User ID: {ev.actor_user_id}")
        assert len(audit_events) >= 2, "Expected at least 2 audit events"
        db.close()
        
        print("\n" + "=" * 70)
        print("        ALL TESTS COMPLETED SUCCESSFULLY! (9/9 PASSED)")
        print("=" * 70)

if __name__ == "__main__":
    run_tests()
