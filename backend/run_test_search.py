import io
import os
import time
from docx import Document as DocxDocument
from fpdf import FPDF
from fastapi.testclient import TestClient

from app.main import app
from app.database import SessionLocal
from app.models.schemas import Document, DocumentPage, DocumentChunk, User, AuditEvent
from app.security import hash_password

def test_hybrid_search():
    with TestClient(app) as client:
        print("=" * 70)
        print("      RAILMIND LITE - DAY 3 HYBRID SEARCH & INDEXING TESTS")
        print("=" * 70)

        total_tests = 9
        passed_tests = 0
    
        db = SessionLocal()
    
        try:
            # Step 0: Ensure docadmin user exists
            print("\n[Init] Ensuring test users exist...")
            docadmin = db.query(User).filter(User.email == "docadmin@camrail.net").first()
            if not docadmin:
                docadmin = User(
                    email="docadmin@camrail.net",
                    password_hash=hash_password("docadminpassword"),
                    full_name="Document Admin",
                    role="document_admin",
                    is_active=True
                )
                db.add(docadmin)
                db.commit()
                db.refresh(docadmin)
    
            # Step 1: Login
            print("\n[Step 1] Authenticating as document_admin...")
            login_res = client.post("/auth/login", data={
                "username": "docadmin@camrail.net",
                "password": "docadminpassword"
            })
            assert login_res.status_code == 200, f"Login failed: {login_res.text}"
            token = login_res.json()["access_token"]
            headers = {"Authorization": f"Bearer {token}"}
            print("-> Login successful! JWT token acquired.")
            passed_tests += 1
    
            # Step 2: Generate test PDF with 2 distinct pages
            print("\n[Step 2] Generating 2-page test PDF (Rest rules on p1, Locomotive maintenance on p2)...")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            pdf.cell(200, 10, text="REGLEMENTATION DU TRAVAIL ET TEMPS DE REPOS", new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.multi_cell(
                0, 10,
                text="Le repos minimum entre deux services est de douze heures. "
                     "Cette regle est strictement appliquee a l ensemble du personnel roulant "
                     "et des agents de conduite de CAMRAIL pour garantir la securite des circulations ferroviaires."
            )
    
            pdf.add_page()
            pdf.cell(200, 10, text="MAINTENANCE DU MATERIEL ROULANT", new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.multi_cell(
                0, 10,
                text="La maintenance preventive des locomotives diesel-electriques s effectue "
                     "tous les cinquante mille kilometres ou tous les six mois au technicentre de Douala Bassa. "
                     "Les bogies et moteurs de traction subissent un controle complet."
            )
    
            pdf_bytes = bytes(pdf.output())
            print(f"-> PDF generated in-memory ({len(pdf_bytes)} bytes)")
            passed_tests += 1
    
            # Step 3: Upload and activate PDF
            print("\n[Step 3] Uploading and activating PDF document...")
            files = {"file": ("consigne_repos_securite.pdf", io.BytesIO(pdf_bytes), "application/pdf")}
            data = {
                "title": "Consigne RH 04 - Temps de Repos et Securite",
                "version": "1.0",
                "category": "Reglementation",
                "department": "Ressources Humaines",
                "effective_date": "2026-08-08"
            }
            upload_res = client.post("/documents", headers=headers, files=files, data=data)
            assert upload_res.status_code == 200, f"Upload failed: {upload_res.text}"
            pdf_doc_id = upload_res.json()["id"]
            assert upload_res.json()["status"] == "indexed", f"Expected status 'indexed', got {upload_res.json()['status']}"
            print(f"-> PDF uploaded and indexed (ID: {pdf_doc_id}, Status: {upload_res.json()['status']})")
    
            activate_res = client.post(f"/documents/{pdf_doc_id}/activate", headers=headers)
            assert activate_res.status_code == 200, f"Activation failed: {activate_res.text}"
            assert activate_res.json()["status"] == "active"
            print(f"-> PDF document {pdf_doc_id} activated successfully!")
            passed_tests += 1
    
            # Step 4: Verify DocumentChunk in database
            print("\n[Step 4] Checking DocumentChunk records in PostgreSQL database...")
            chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == pdf_doc_id).order_by(DocumentChunk.chunk_index).all()
            assert len(chunks) >= 2, f"Expected at least 2 chunks, found {len(chunks)}"
            for c in chunks:
                assert c.embedding is not None, f"Chunk {c.id} has null embedding"
                assert c.search_vector is not None, f"Chunk {c.id} has null search_vector"
                print(f"   * Chunk {c.chunk_index}: Page {c.page_start}, length={len(c.content)} chars, embedding=Vector(384)")
            print(f"-> DocumentChunk verification passed! ({len(chunks)} chunks with valid embeddings and search vectors)")
            passed_tests += 1
    
            # Step 5: Semantic search for target phrase
            print("\n[Step 5] Testing Hybrid Search with semantic query: 'durée minimale de repos entre deux services'...")
            start_time = time.time()
            search_res = client.post("/search", headers=headers, json={
                "query": "durée minimale de repos entre deux services",
                "top_k": 5
            })
            elapsed = time.time() - start_time
            assert search_res.status_code == 200, f"Search failed: {search_res.text}"
            search_data = search_res.json()
            assert len(search_data["results"]) > 0, "No results returned for relevant query!"
            
            top_result = search_data["results"][0]
            print(f"-> Top search result (latency: {elapsed:.3f}s):")
            print(f"   * Document: {top_result['document_title']}")
            print(f"   * Page Start: {top_result['page_start']}")
            print(f"   * Score RRF: {top_result['score']}")
            print(f"   * Excerpt: {top_result['excerpt']}")
            
            assert top_result["page_start"] == 1, f"Expected result from Page 1 (Rest rules), got Page {top_result['page_start']}"
            assert "repos" in top_result["excerpt"].lower()
            print("-> Success! Hybrid search ranked Page 1 (rest rules) as #1 top result!")
            passed_tests += 1
    
            # Step 6: Testing Hybrid Search with lexical focus
            print(f"\n[Step 6] Testing Hybrid Search with lexical query: 'technicentre de Douala'...")
            search_res_lex = client.post(
                "/search",
                headers=headers,
                json={"query": "technicentre de Douala", "top_k": 5}
            )
            assert search_res_lex.status_code == 200
            results_lex = search_res_lex.json()["results"]
            assert len(results_lex) > 0
            top_result_lex = results_lex[0]
            assert top_result_lex["page_start"] == 2, "Lexical search should rank Page 2 (maintenance) higher!"
            print("-> Success! Lexical-heavy search ranked Page 2 as #1 top result!")
            passed_tests += 1
            # Step 6.5: Testing Hybrid Search with partial natural language match
            print(f"\n[Step 6.5] Testing Hybrid Search with partial match: 'quelle est la pause obligatoire entre deux services'...")
            search_res_partial = client.post(
                "/search",
                headers=headers,
                json={"query": "quelle est la pause obligatoire entre deux services", "top_k": 5}
            )
            assert search_res_partial.status_code == 200
            results_partial = search_res_partial.json()["results"]
            assert len(results_partial) > 0, "Should return results despite partial lexical match"
            
            # Verify that the document about rest (Page 1) is present in the results
            page_starts = [r["page_start"] for r in results_partial]
            assert 1 in page_starts, "Page 1 (rest rules) should be found by partial lexical OR vector search!"
            print("-> Success! Partial natural language match found the relevant document.")
            passed_tests += 1
    
    
            # Step 7: DOCX upload, activation and citation check
            print("\n[Step 7] Generating, uploading and searching a DOCX document...")
            docx_file = DocxDocument()
            docx_file.add_heading("Consigne Exploitation - Cantonnement Telephonique", 0)
            docx_file.add_paragraph(
                "Procedure operationnelle pour le cantonnement telephonique en cas de derangement des signaux "
                "automatiques sur la section ferroviaire Douala Yaounde. "
                "Les chefs de securite doivent echanger les depeches reglementaires."
            )
            docx_io = io.BytesIO()
            docx_file.save(docx_io)
            docx_bytes = docx_io.getvalue()
    
            docx_files = {"file": ("consigne_cantonnement.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            docx_data = {
                "title": "Consigne Exploitation Cantonnement",
                "version": "2.1",
                "category": "Exploitation",
                "department": "Securite Ferroviaire",
                "effective_date": "2026-08-08"
            }
            docx_upload_res = client.post("/documents", headers=headers, files=docx_files, data=docx_data)
            assert docx_upload_res.status_code == 200
            docx_doc_id = docx_upload_res.json()["id"]
    
            client.post(f"/documents/{docx_doc_id}/activate", headers=headers)
            print(f"-> DOCX uploaded and activated (ID: {docx_doc_id})")
    
            docx_search_res = client.post("/search", headers=headers, json={
                "query": "cantonnement telephonique derangement des signaux",
                "top_k": 5
            })
            assert docx_search_res.status_code == 200
            docx_search_data = docx_search_res.json()
            assert len(docx_search_data["results"]) > 0
    
            # Check for citation format
            matching_docx_result = next((r for r in docx_search_data["results"] if r["document_id"] == docx_doc_id), None)
            assert matching_docx_result is not None, "DOCX chunk not found in search results"
            assert matching_docx_result["is_full_document_citation"] is True, f"Expected is_full_document_citation=True for DOCX, got {matching_docx_result['is_full_document_citation']}"
            print(f"-> DOCX search result confirmed with is_full_document_citation=True: {matching_docx_result['is_full_document_citation']}")
            passed_tests += 1
    
            # Step 8: Clean up test documents
            print("\n[Step 8] Cleaning up test documents and files...")
            pdf_record = db.query(Document).filter(Document.id == pdf_doc_id).first()
            docx_record = db.query(Document).filter(Document.id == docx_doc_id).first()
    
            for rec in [pdf_record, docx_record]:
                if rec and rec.file_url and os.path.exists(rec.file_url):
                    try:
                        os.remove(rec.file_url)
                    except Exception:
                        pass
    
            print("-> Storage cleanup completed!")
            passed_tests += 1
    
            print("\n" + "=" * 70)
            print(f"        ALL TESTS COMPLETED SUCCESSFULLY! ({passed_tests}/{total_tests} PASSED)")
            print("=" * 70)
    
        finally:
            db.close()


if __name__ == "__main__":
    test_hybrid_search()
